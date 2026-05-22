import subprocess
import sys
import os
import re
from io import BytesIO

import streamlit as st
from docx import Document

# Auto-install Playwright browser binaries if missing
try:
    from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError, Error as PlaywrightError
except ImportError:
    st.error("📦 Installing Playwright browser binaries... This may take a moment.")
    st.info("Please refresh the page in a few seconds after installation completes.")
    try:
        subprocess.check_call([sys.executable, "-m", "playwright", "install", "chromium"])
        st.rerun()
    except Exception as install_error:
        st.error(f"Failed to install Playwright: {install_error}")
        st.stop()
    from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError, Error as PlaywrightError

st.set_page_config(
    page_title="ChatGPT Share → PDF",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# Custom CSS for better styling
st.markdown(
    """
    <style>
    /* Main container styling */
    .main {
        padding-top: 1rem;
    }
    
    /* Button styling */
    .stButton > button {
        border-radius: 8px;
        font-weight: 600;
        padding: 0.75rem 2rem;
        font-size: 1rem;
    }
    
    /* Input styling */
    .stTextInput > div > div > input {
        border-radius: 8px;
    }
    
    /* Subheader styling */
    .stSubheader {
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    
    /* Success/Error messages */
    .stAlert {
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Progress bar styling */
    .stProgress > div > div > div {
        background-color: #10a37f;
        border-radius: 4px;
    }

    /* Primary PDF button */
    .stButton > button[kind="primary"] {
        background-color: #dc2626 !important;
        border-color: #dc2626 !important;
        color: #ffffff !important;
    }
    .stButton > button[kind="primary"]:hover {
        background-color: #b91c1c !important;
        border-color: #b91c1c !important;
    }

    /* Secondary Word button */
    .stButton > button[kind="secondary"] {
        background-color: #2563eb !important;
        border-color: #2563eb !important;
        color: #ffffff !important;
    }
    .stButton > button[kind="secondary"]:hover {
        background-color: #1d4ed8 !important;
        border-color: #1d4ed8 !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

TITLE = "📄 ChatGPT Share Link Converter"
DESCRIPTION = (
    "Paste any public ChatGPT share link and convert it to a clean PDF or Word document. "
    "This app uses Playwright to render the page exactly as the browser does before exporting it."
)

CSS_CLEANUP = """
header, nav, footer, [role='banner'], [role='navigation'], .flex.flex-col.items-center, .prose a[href*='login'], .overflow-hidden, .button, .btn, .sticky, [data-testid='sidebar'] { display: none !important; }
body { background: #ffffff !important; color: #111827 !important; }
"""

VALID_URL_PATTERN = re.compile(
    r"^https?://(?:www\.)?(?:chat\.openai\.com/share/|chatgpt\.com/share/|shareg\.pt/)[\w\-]+(?:\?.*)?$",
    re.IGNORECASE,
)


def is_chatgpt_share_url(url: str) -> bool:
    return bool(VALID_URL_PATTERN.match(url.strip()))


@st.cache_resource
def verify_playwright_installation():
    """Verify that Playwright browser binaries are installed."""
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as pw:
            # Try to launch to verify binaries exist
            browser = pw.chromium.launch(headless=True)
            browser.close()
        return True
    except Exception as e:
        error_msg = str(e)
        if "Executable doesn't exist" in error_msg or "ENOENT" in error_msg:
            return False
        # Other errors, still return True to attempt the conversion
        return True


def export_chatgpt_share_to_pdf(url: str) -> bytes:
    with sync_playwright() as playwright:
        try:
            browser = playwright.chromium.launch(headless=True)
        except PlaywrightError as exc:
            if "Executable doesn't exist" in str(exc) or "playwright install" in str(exc).lower():
                raise RuntimeError(
                    "Playwright browser binaries are missing. "
                    "Run `python -m playwright install chromium` and restart the app."
                ) from exc
            raise

        context = browser.new_context(viewport={"width": 1200, "height": 1400})
        page = context.new_page()

        page.goto(url, wait_until="networkidle", timeout=60000)
        page.wait_for_timeout(1200)

        # Wait for the chat content to appear before exporting.
        page.wait_for_selector("main, article, [data-testid='share-view']", timeout=45000)
        page.add_style_tag(content=CSS_CLEANUP)
        page.wait_for_timeout(800)

        pdf_bytes = page.pdf(
            format="A4",
            print_background=True,
            margin={"top": "0.5in", "bottom": "0.5in", "left": "0.5in", "right": "0.5in"},
        )

        context.close()
        browser.close()

    return pdf_bytes


def export_chatgpt_share_to_docx(url: str) -> bytes:
    with sync_playwright() as playwright:
        try:
            browser = playwright.chromium.launch(headless=True)
        except PlaywrightError as exc:
            if "Executable doesn't exist" in str(exc) or "playwright install" in str(exc).lower():
                raise RuntimeError(
                    "Playwright browser binaries are missing. "
                    "Run `python -m playwright install chromium` and restart the app."
                ) from exc
            raise

        context = browser.new_context(viewport={"width": 1200, "height": 1400})
        page = context.new_page()

        page.goto(url, wait_until="networkidle", timeout=60000)
        page.wait_for_timeout(1200)

        page.wait_for_selector("main, article, [data-testid='share-view']", timeout=45000)
        page.wait_for_timeout(800)

        chat_text = page.inner_text("main, article, [data-testid='share-view']")
        document = Document()
        document.add_heading("ChatGPT Share Export", level=1)
        document.add_paragraph(f"Source: {url}")
        document.add_paragraph("")

        for line in chat_text.splitlines():
            if not line.strip():
                document.add_paragraph("")
            else:
                document.add_paragraph(line)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        context.close()
        browser.close()

    return buffer.read()


def main() -> None:
    # Verify Playwright installation
    if not verify_playwright_installation():
        st.error(
            "❌ **Playwright browser binaries are missing!**\n\n"
            "This error should resolve automatically on Streamlit Cloud within a few minutes. "
            "Please refresh the page to retry.\n\n"
            "If the issue persists, the app is being re-initialized. Try refreshing again."
        )
        st.info("💡 **This is a one-time setup issue** - it typically resolves itself automatically.")
        st.stop()
    
    # Sidebar information
    with st.sidebar:
        st.markdown("### ℹ️ About")
        st.write(
            "This tool converts public ChatGPT share links into clean PDF or Word documents."
        )
        
        st.markdown("### 🎯 Supported Links")
        st.code(
            "chatgpt.com/share/...\n"
            "chat.openai.com/share/...\n"
            "shareg.pt/...",
            language="text",
        )
        
        st.markdown("### ⚙️ Features")
        st.markdown(
            "✓ Full JavaScript rendering\n"
            "✓ Clean formatting\n"
            "✓ High-quality PDF export\n"
            "✓ Public links only"
        )
        
        st.markdown("---")
        st.caption("Made with ❤️ using Streamlit + Playwright by Surajit00007")
    
    # Page header
    col1, col2 = st.columns([1, 4])
    with col1:
        st.markdown("# 📄", unsafe_allow_html=True)
    with col2:
        st.title("ChatGPT Share Link Converter")
    
    st.markdown("---")
    st.markdown(
        "<p style='text-align: center; color: #666;'>"
        "Convert any public ChatGPT share link into a clean, downloadable PDF or Word document"
        "</p>",
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # Initialize session state
    if "pdf_data" not in st.session_state:
        st.session_state.pdf_data = None
    if "conversion_complete" not in st.session_state:
        st.session_state.conversion_complete = False

    # Input section
    st.subheader("🔗 Enter your ChatGPT share link")
    share_url = st.text_input(
        "Paste your public ChatGPT share link:",
        placeholder="https://chatgpt.com/share/abc123...",
        help="Supports: chatgpt.com/share, chat.openai.com/share, shareg.pt",
    )

    # URL validation feedback
    if share_url:
        if is_chatgpt_share_url(share_url):
            st.success("✓ Valid ChatGPT share link")
        else:
            st.error(
                "✗ Invalid link format. Please use:\n"
                "- `chatgpt.com/share/...`\n"
                "- `chat.openai.com/share/...`\n"
                "- `shareg.pt/...`"
            )
            return

    if not share_url:
        st.info("👉 Paste a ChatGPT share link above to get started", icon="ℹ️")
        return

    # Conversion section
    st.subheader("📥 Convert the chat share link")
    
    col1, col2 = st.columns(2)
    
    with col1:
        convert_pdf = st.button(
            "🚀 Convert to PDF",
            use_container_width=True,
            type="primary",
        )
    with col2:
        convert_word = st.button(
            "📝 Convert to Word",
            use_container_width=True,
            type="secondary",
        )

    if convert_pdf or convert_word:
        with st.spinner(""):
            progress_placeholder = st.empty()
            status_placeholder = st.empty()

            try:
                progress_placeholder.progress(0, text="⏳ Launching browser...")
                status_placeholder.write("Starting Chromium headless browser")

                if convert_pdf:
                    result_bytes = export_chatgpt_share_to_pdf(share_url)
                    st.session_state.pdf_data = result_bytes
                    progress_placeholder.progress(100, text="✓ PDF generated")
                    status_placeholder.empty()
                    st.success("✅ PDF generated successfully! Ready to download.", icon="✅")

                if convert_word:
                    result_bytes = export_chatgpt_share_to_docx(share_url)
                    st.session_state.docx_data = result_bytes
                    progress_placeholder.progress(100, text="✓ Word document generated")
                    status_placeholder.empty()
                    st.success("✅ Word document generated successfully! Ready to download.", icon="✅")

            except PlaywrightTimeoutError:
                progress_placeholder.empty()
                status_placeholder.empty()
                st.error(
                    "⏱️ Timeout: The page took too long to render.\n\n"
                    "**Try:**\n"
                    "- Verify the link is public\n"
                    "- Try again in a few moments\n"
                    "- Check if the conversation is very long"
                )
            except Exception as exc:
                progress_placeholder.empty()
                status_placeholder.empty()
                st.error(f"❌ Error: {exc}")

    # Download section
    if st.session_state.pdf_data or st.session_state.get("docx_data"):
        st.markdown("---")
        st.subheader("📥 Download your generated file")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.session_state.pdf_data:
                st.download_button(
                    "⬇️ Download PDF",
                    st.session_state.pdf_data,
                    file_name="chatgpt-share.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                )
        with col2:
            if st.session_state.get("docx_data"):
                st.download_button(
                    "⬇️ Download Word",
                    st.session_state.docx_data,
                    file_name="chatgpt-share.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        file_size_lines = []
        if st.session_state.pdf_data:
            file_size_lines.append(f"PDF size: {len(st.session_state.pdf_data) / 1024:.1f} KB")
        if st.session_state.get("docx_data"):
            file_size_lines.append(f"Word size: {len(st.session_state.docx_data) / 1024:.1f} KB")
        if file_size_lines:
            st.caption(" | ".join(file_size_lines))


if __name__ == "__main__":
    main()
